import os
import json
import uuid
import fitz  # PyMuPDF
import requests
from flask import Flask, request, jsonify, send_file, render_template
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10MB max

OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY", "")
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL_CANDIDATES = [
    "openrouter/free",
    "meta-llama/llama-3.3-70b-instruct:free",
    "openai/gpt-oss-20b:free",
    "openai/gpt-oss-120b:free",
    "nvidia/nemotron-3-nano-9b-v2:free",
]

COMPANY_PRESETS = {
    "google": "Google-style resume: clean single page, ATS-friendly, no photo. Sections: Summary (2 lines), Skills (grouped by category), Experience (reverse chronological, 3-4 bullet points per role using STAR format with metrics), Education, Projects.",
    "amazon": "Amazon Leadership Principles resume: results-oriented, heavy on quantifiable metrics. Sections: Summary, Core Skills, Experience (STAR-format bullets showing Ownership and Deliver Results), Projects, Education.",
    "microsoft": "Microsoft-style: technical focus. Sections: Summary, Technical Skills (languages, tools, platforms), Work Experience with impact metrics, Projects with tech stack, Education, Certifications.",
    "mckinsey": "McKinsey consulting style: 1 page strictly. Bullet points with strong action verbs, quantified impact. Sections: Education (at top), Experience, Leadership & Activities. No objective/summary.",
    "goldman": "Goldman Sachs finance resume: very clean and concise. Education first with GPA. Experience with deal/project values. Skills section. All dates right-aligned. Max 1 page.",
    "startup": "Modern startup resume: brief 2-line summary, skills with proficiency. Experience focused on growth and ownership. Side projects highlighted. Clean modern layout.",
}

# ─── Section key mapping ─────────────────────────────────────────────────────
# Maps common section name variations to internal keys used in build_docx()
SECTION_KEY_MAP = {
    "summary": "summary", "objective": "summary", "profile": "summary",
    "about": "summary", "professional summary": "summary",
    "skills": "skills", "technical skills": "skills", "core competencies": "skills",
    "competencies": "skills", "technologies": "skills", "expertise": "skills",
    "experience": "experience", "work experience": "experience",
    "work history": "experience", "professional experience": "experience",
    "employment": "experience", "career history": "experience",
    "education": "education", "academic background": "education",
    "qualifications": "education",
    "projects": "projects", "key projects": "projects", "personal projects": "projects",
    "certifications": "certifications", "achievements": "certifications",
    "certifications & achievements": "certifications", "awards": "certifications",
}


def normalize_section_key(name):
    return SECTION_KEY_MAP.get(name.lower().strip(), None)


def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text.strip()


def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return text.strip()


# ─── Regex contact extraction ─────────────────────────────────────────────────
_EMAIL_REGEX = re.compile(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+')
_PHONE_REGEX = re.compile(r'(\+?\d{1,3}[-.\s]?)?\d{10}')
_LINKEDIN_REGEX = re.compile(r'linkedin\.com/in/[A-Za-z0-9\-_/]+', re.IGNORECASE)
_GITHUB_REGEX = re.compile(r'github\.com/[A-Za-z0-9\-_/]+', re.IGNORECASE)


def extract_contact_info(resume_text):
    contact = {}
    email_match = _EMAIL_REGEX.search(resume_text)
    if email_match:
        contact["email"] = email_match.group(0)
    phone_match = _PHONE_REGEX.search(resume_text)
    if phone_match:
        contact["phone"] = phone_match.group(0).strip()
    linkedin_match = _LINKEDIN_REGEX.search(resume_text)
    if linkedin_match:
        contact["linkedin"] = linkedin_match.group(0)
    github_match = _GITHUB_REGEX.search(resume_text)
    if github_match:
        contact["github"] = github_match.group(0)
    if contact.get("email"):
        for line in resume_text.split("\n"):
            if contact["email"] not in line:
                continue
            for part in re.split(r"[·|]", line):
                part = part.strip().strip(",")
                if not part:
                    continue
                if contact["email"] in part:
                    continue
                if contact.get("phone") and contact["phone"] in part:
                    continue
                if "," in part or (len(part) < 40 and not any(ch.isdigit() for ch in part)):
                    contact["address"] = part
                    break
            break
    return contact


# ─── Template structure analyzer ─────────────────────────────────────────────
def analyze_template_structure(template_text):
    """
    Calls AI to extract structural metadata from a company template:
    section order, exact section names, bullet prefix, name alignment.
    Returns a dict like:
    {
      "section_order": ["education", "skills", "experience", "projects", "certifications"],
      "section_names": {
        "education": "EDUCATION",
        "skills": "[ CORE COMPETENCIES ]",
        "experience": "WORK HISTORY",
        "projects": "KEY PROJECTS",
        "certifications": "ACHIEVEMENTS"
      },
      "bullet_prefix": ">> ",
      "name_alignment": "left"
    }
    """
    system_prompt = """You are a resume format analyzer. Given a sample resume, extract its structural metadata.

Return ONLY a JSON object with these exact keys:
{
  "section_order": ["list of section keys in the order they appear — use only: summary, skills, experience, education, projects, certifications"],
  "section_names": {
    "summary": "exact heading text as written in the template",
    "skills": "exact heading text",
    "experience": "exact heading text",
    "education": "exact heading text",
    "projects": "exact heading text",
    "certifications": "exact heading text"
  },
  "bullet_prefix": "the bullet prefix used (e.g. '>> ', '• ', '- ', '▸ ') — just the prefix string",
  "name_alignment": "left or center"
}

Only include sections that actually exist in the template in section_order and section_names.
Return ONLY valid JSON. No explanation, no markdown."""

    user_prompt = f"Analyze this resume template and extract its structure:\n\n{template_text}"

    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost:5000",
    }

    last_error = None
    for model in MODEL_CANDIDATES:
        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "temperature": 0.1,
            "max_tokens": 500,
        }
        try:
            response = requests.post(OPENROUTER_URL, headers=headers, json=payload, timeout=30)
            response.raise_for_status()
        except requests.exceptions.HTTPError as e:
            print(f"[analyze_template] Model '{model}' failed: {response.text}")
            last_error = e
            continue
        except requests.exceptions.RequestException as e:
            last_error = e
            continue

        content = response.json()["choices"][0]["message"]["content"]
        if not content:
            continue
        content = content.strip()
        content = re.sub(r"^```json\s*", "", content)
        content = re.sub(r"\s*```$", "", content)
        content = re.sub(r",\s*}", "}", content)
        content = re.sub(r",\s*]", "]", content)

        try:
            result = json.loads(content)
            print(f"[analyze_template] Succeeded: {json.dumps(result, indent=2)}")
            return result
        except json.JSONDecodeError:
            continue

    # fallback: return default structure if analysis fails
    print("[analyze_template] Failed, using default structure")
    return {
        "section_order": ["summary", "skills", "experience", "projects", "education", "certifications"],
        "section_names": {},
        "bullet_prefix": "• ",
        "name_alignment": "center"
    }


def call_ai(resume_text, format_description):
    system_prompt = """You are a professional resume writer. You receive a candidate's resume text and formatting instructions.
Your job: rewrite and reformat the resume exactly according to the company format described.

IMPORTANT: Contact details (email, phone, address, LinkedIn, GitHub) must ALWAYS be included if they exist anywhere in the original resume text, even if formatted with icons, symbols, or unusual spacing. Never omit them.

IMPORTANT: Include ALL education entries found in the original resume — undergraduate degree, intermediate/12th, high school/10th, diplomas, etc. Do not drop earlier schooling. Each entry must appear in the "education" array.

Return the reformatted resume in a structured JSON format with these keys:
{
  "name": "Full Name",
  "contact": {
    "email": "email@example.com",
    "phone": "+91 XXXXXXXXXX",
    "address": "City, State",
    "linkedin": "linkedin.com/in/username",
    "github": "github.com/username"
  },
  "summary": "2-3 line summary",
  "skills": [{"category": "Languages", "items": ["Python", "Java"]}],
  "experience": [{"title": "Job Title", "company": "Company", "duration": "Jan 2022 - Dec 2023", "bullets": ["bullet 1", "bullet 2"]}],
  "education": [{"degree": "B.Tech CSE", "institution": "SRM University", "year": "2022-2026", "gpa": "8.7"}],
  "projects": [{"name": "Project Name", "tech": "Python, Flask", "bullets": ["bullet 1"]}],
  "certifications": ["Cert 1", "Cert 2"]
}
Only omit a contact field if it genuinely does not exist. Keep all real info. Do not make up data."""

    user_prompt = f"""Company format requirements:
{format_description}

Candidate's resume:
{resume_text}

Return ONLY valid JSON. No explanation, no markdown, no code blocks."""

    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost:5000",
    }

    last_error = None
    for model in MODEL_CANDIDATES:
        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "temperature": 0.3,
            "max_tokens": 3000,
        }
        try:
            response = requests.post(OPENROUTER_URL, headers=headers, json=payload, timeout=60)
            response.raise_for_status()
        except requests.exceptions.HTTPError as e:
            print(f"[call_ai] Model '{model}' failed ({response.status_code}): {response.text}")
            last_error = requests.exceptions.HTTPError(f"{e} | OpenRouter says: {response.text}")
            continue
        except requests.exceptions.RequestException as e:
            print(f"[call_ai] Model '{model}' request error: {e}")
            last_error = e
            continue

        content = response.json()["choices"][0]["message"]["content"]
        if not content:
            print(f"[call_ai] Model '{model}' returned empty content, trying next.")
            last_error = Exception(f"Model '{model}' returned empty content.")
            continue
        content = content.strip()
        content = re.sub(r"^```json\s*", "", content)
        content = re.sub(r"\s*```$", "", content)
        content = re.sub(r",\s*}", "}", content)
        content = re.sub(r",\s*]", "]", content)
        content = content.strip()

        try:
            parsed = json.loads(content)
            print(f"[call_ai] Succeeded using model: {model}")
            return parsed
        except json.JSONDecodeError as e:
            print(f"[call_ai] Model '{model}' returned invalid JSON: {e}\nRaw: {content[:300]}")
            last_error = Exception(f"AI returned invalid JSON: {e}")
            continue

    raise last_error if last_error else Exception("All AI model candidates failed.")


def build_docx(data, company_name, template_style=None):
    print(f"[build_docx] template_style received: {json.dumps(template_style, indent=2) if template_style else 'None'}")
    print(f"[build_docx] section_order will be: {template_style.get('section_order') if template_style else 'DEFAULT'}")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Resolve template style or use defaults
    if template_style:
        section_order  = template_style.get("section_order", ["summary","skills","experience","projects","education","certifications"])
        section_names  = template_style.get("section_names", {})
        bullet_prefix  = template_style.get("bullet_prefix", "• ")
        name_alignment = template_style.get("name_alignment", "center")
    else:
        section_order  = ["summary","skills","experience","projects","education","certifications"]
        section_names  = {}
        bullet_prefix  = ""       # empty = use List Bullet style
        name_alignment = "center"

    def get_heading_text(key, fallback):
        return section_names.get(key, fallback)

    def add_heading(text):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1a1a2e')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def add_bullet(text):
        if bullet_prefix:
            # Custom prefix — plain paragraph with prefix text
            p = doc.add_paragraph()
            run = p.add_run(f"{bullet_prefix}{text}")
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.left_indent = Inches(0.25)
        else:
            # Default Word bullet style
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(text)
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.left_indent = Inches(0.25)
        return p

    # --- NAME ---
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT if name_alignment == "left" else WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(data.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    name_para.paragraph_format.space_after = Pt(2)

    # --- CONTACT ---
    contact = data.get("contact")
    if contact:
        if isinstance(contact, dict):
            parts = [contact.get(k) for k in ("email","phone","address","linkedin","github") if contact.get(k)]
            contact_text = "  |  ".join(parts)
        else:
            contact_text = str(contact)
        if contact_text:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT if name_alignment == "left" else WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(contact_text)
            contact_run.font.size = Pt(9)
            contact_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            contact_para.paragraph_format.space_after = Pt(6)

    # --- SECTIONS in template order ---
    def render_summary():
        if data.get("summary"):
            add_heading(get_heading_text("summary", "SUMMARY"))
            p = doc.add_paragraph(data["summary"])
            p.runs[0].font.size = Pt(10)
            p.paragraph_format.space_after = Pt(4)

    def render_skills():
        if data.get("skills"):
            add_heading(get_heading_text("skills", "SKILLS"))
            for skill_group in data["skills"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                cat_run = p.add_run(skill_group.get("category", "") + ": ")
                cat_run.bold = True
                cat_run.font.size = Pt(10)
                items_run = p.add_run(", ".join(skill_group.get("items", [])))
                items_run.font.size = Pt(10)

    def render_experience():
        if data.get("experience"):
            add_heading(get_heading_text("experience", "EXPERIENCE"))
            for exp in data["experience"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                title_run = p.add_run(exp.get("title", ""))
                title_run.bold = True
                title_run.font.size = Pt(10.5)
                company_run = p.add_run(f"  |  {exp.get('company', '')}")
                company_run.font.size = Pt(10)
                company_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                if exp.get("duration"):
                    dur_run = p.add_run(f"  —  {exp['duration']}")
                    dur_run.font.size = Pt(9)
                    dur_run.italic = True
                    dur_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                for bullet in exp.get("bullets", []):
                    add_bullet(bullet)

    def render_projects():
        if data.get("projects"):
            add_heading(get_heading_text("projects", "PROJECTS"))
            for proj in data["projects"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                name_run = p.add_run(proj.get("name", ""))
                name_run.bold = True
                name_run.font.size = Pt(10.5)
                if proj.get("tech"):
                    tech_run = p.add_run(f"  |  {proj['tech']}")
                    tech_run.font.size = Pt(9)
                    tech_run.italic = True
                    tech_run.font.color.rgb = RGBColor(0x55, 0x55, 0x88)
                for bullet in proj.get("bullets", []):
                    add_bullet(bullet)

    def render_education():
        if data.get("education"):
            add_heading(get_heading_text("education", "EDUCATION"))
            for edu in data["education"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                deg_run = p.add_run(edu.get("degree", ""))
                deg_run.bold = True
                deg_run.font.size = Pt(10.5)
                inst_run = p.add_run(f"  |  {edu.get('institution', '')}")
                inst_run.font.size = Pt(10)
                if edu.get("year"):
                    yr_run = p.add_run(f"  —  {edu['year']}")
                    yr_run.font.size = Pt(9)
                    yr_run.italic = True
                    yr_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                if edu.get("gpa"):
                    p2 = doc.add_paragraph()
                    p2.paragraph_format.space_after = Pt(1)
                    p2.paragraph_format.left_indent = Inches(0.2)
                    gpa_run = p2.add_run(f"GPA: {edu['gpa']}")
                    gpa_run.font.size = Pt(9)
                    gpa_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    def render_certifications():
        if data.get("certifications"):
            add_heading(get_heading_text("certifications", "CERTIFICATIONS"))
            for cert in data["certifications"]:
                add_bullet(cert)

    render_map = {
        "summary": render_summary,
        "skills": render_skills,
        "experience": render_experience,
        "projects": render_projects,
        "education": render_education,
        "certifications": render_certifications,
    }

    # Render sections in template order
    rendered = set()
    for key in section_order:
        if key in render_map:
            render_map[key]()
            rendered.add(key)

    # Render any remaining sections not in template order (fallback)
    for key, fn in render_map.items():
        if key not in rendered:
            fn()

    os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)
    filename = f"resume_{company_name}_{uuid.uuid4().hex[:6]}.docx"
    filepath = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    doc.save(filepath)
    return filepath, filename


# ─── ROUTES ──────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html", presets=list(COMPANY_PRESETS.keys()))


@app.route("/convert", methods=["POST"])
def convert():
    if "resume" not in request.files:
        return jsonify({"error": "No resume file uploaded"}), 400

    file = request.files["resume"]
    company = request.form.get("company", "google").lower()
    custom_format = request.form.get("custom_format", "").strip()
    template_file = request.files.get("company_template")

    if not (file.filename.lower().endswith(".pdf") or file.filename.lower().endswith(".docx")):
        return jsonify({"error": "Only PDF or DOCX files are supported"}), 400

    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    ext = "pdf" if file.filename.lower().endswith(".pdf") else "docx"
    resume_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{uuid.uuid4().hex}.{ext}")
    file.save(resume_path)

    template_path = None
    template_style = None

    try:
        if ext == "pdf":
            resume_text = extract_text_from_pdf(resume_path)
        else:
            resume_text = extract_text_from_docx(resume_path)

        if not resume_text:
            return jsonify({"error": "Could not extract text from resume"}), 400

        print("===== EXTRACTED RESUME TEXT =====")
        print(resume_text[:500])
        print("==================================")

        if template_file and template_file.filename.lower().endswith(".docx"):
            template_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{uuid.uuid4().hex}.docx")
            template_file.save(template_path)

            try:
                template_text = extract_text_from_docx(template_path)
            except Exception as e:
                return jsonify({"error": f"Could not read template DOCX: {str(e)}"}), 400

            # Analyze template structure FIRST (section order, names, bullet style)
            template_style = analyze_template_structure(template_text)

            format_desc = (
                "Use the following resume as a STRUCTURAL AND STYLE reference. "
                "Mimic its section order, section naming, and level of detail/tone. "
                "Do NOT copy its actual content — only its format and style.\n\n"
                f"Reference resume:\n{template_text}"
            )
        elif custom_format:
            format_desc = custom_format
        else:
            format_desc = COMPANY_PRESETS.get(company, COMPANY_PRESETS["google"])

        structured_data = call_ai(resume_text, format_desc)

        # Regex contact override
        regex_contact = extract_contact_info(resume_text)
        ai_contact = structured_data.get("contact")
        if not isinstance(ai_contact, dict):
            ai_contact = {}
        structured_data["contact"] = {**ai_contact, **regex_contact}

        print("===== FINAL CONTACT =====")
        print(structured_data["contact"])
        print("===== TEMPLATE STYLE =====")
        print(json.dumps(template_style, indent=2) if template_style else "None (preset/custom mode)")
        print("===== AI OUTPUT =====")
        print(json.dumps(structured_data, indent=2))

        filepath, filename = build_docx(structured_data, company, template_style=template_style)

        return jsonify({"success": True, "filename": filename, "preview": structured_data, "template_style": template_style})

    except json.JSONDecodeError as e:
        return jsonify({"error": f"AI returned invalid JSON: {str(e)}"}), 500
    except requests.RequestException as e:
        return jsonify({"error": f"API call failed: {str(e)}"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if os.path.exists(resume_path):
            os.remove(resume_path)
        if template_path and os.path.exists(template_path):
            os.remove(template_path)


@app.route("/download/<filename>")
def download(filename):
    filepath = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    if not os.path.exists(filepath):
        return jsonify({"error": "File not found"}), 404
    return send_file(filepath, as_attachment=True, download_name=filename)


if __name__ == "__main__":
    app.run(debug=True, port=5000)